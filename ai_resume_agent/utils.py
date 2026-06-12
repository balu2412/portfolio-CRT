import io
from pypdf import PdfReader
import docx

def extract_text_from_pdf(pdf_file) -> str:
    """Extracts text from a PDF file-like object or file path."""
    try:
        reader = PdfReader(pdf_file)
        text = ""
        for page in reader.pages:
            content = page.extract_text()
            if content:
                text += content + "\n"
        return text.strip()
    except Exception as e:
        return f"Error reading PDF: {str(e)}"

def extract_text_from_docx(docx_file) -> str:
    """Extracts text from a DOCX file-like object or file path."""
    try:
        # docx.Document needs a file path or a stream that supports seek.
        # If it's bytes or a file-like object, wrap it in BytesIO if needed.
        if isinstance(docx_file, bytes):
            doc = docx.Document(io.BytesIO(docx_file))
        else:
            doc = docx.Document(docx_file)
        
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text.append(cell.text)
        return "\n".join(text).strip()
    except Exception as e:
        return f"Error reading DOCX: {str(e)}"

def get_sample_resume() -> str:
    return """ALEX MERCER
Email: alex.mercer@email.dev | Phone: +1 (555) 019-2834 | GitHub: github.com/alexmercer
Portfolio: alexmercer.dev | Location: San Francisco, CA

PROFESSIONAL SUMMARY
Senior Software Engineer with 6+ years of experience specializing in Full-Stack Web Development, LLMs, and Agentic Systems. Proven track record of designing scalable APIs, crafting interactive React interfaces, and orchestrating multi-agent LLM systems using LangChain and LangGraph.

TECHNICAL SKILLS
- Languages: Python, TypeScript, JavaScript, SQL, HTML/CSS
- Frameworks: React, Next.js, FastAPI, Django, Express.js
- AI/LLM: LangChain, LangGraph, OpenAI API, Gemini API, RAG (Retrieval-Augmented Generation), FAISS
- Databases & Devops: PostgreSQL, Redis, MongoDB, Docker, AWS (S3, EC2, RDS), Git, CI/CD

PROFESSIONAL EXPERIENCE
Senior Full-Stack Engineer | TechCorp Inc.
Jan 2022 - Present | San Francisco, CA
- Led a team of 3 developers to rebuild the core customer engagement platform using Next.js and FastAPI, improving page load speeds by 40% and database throughput by 25%.
- Architected and deployed a multi-agent customer support routing system using LangGraph and LangChain, reducing average ticket response times from 4 hours to under 2 minutes.
- Implemented vector search capabilities using FAISS and OpenAI embeddings to enable context-aware semantic search over 100k+ support documents.
- Standardized Docker-based local development environments and set up automated CI/CD pipelines on AWS (ECS, CodePipeline).

Full-Stack Developer | InnovateSoft Ltd.
Jun 2020 - Dec 2021 | Austin, TX
- Developed and maintained responsive web applications using React, Redux, and Node.js.
- Designed and optimized complex SQL queries and database schemas in PostgreSQL, reducing query latency by 15%.
- Created custom RESTful APIs using Django Rest Framework (DRF) to support mobile and web clients.
- Maintained cloud assets on AWS S3 and EC2, implementing automated backup strategies.

EDUCATION
Master of Science in Computer Science
State University | Graduated: May 2020

Bachelor of Science in Computer Engineering
Texas Tech College | Graduated: May 2018

KEY PROJECTS
AI-Powered Resume Screen & Prep Platform (LangGraph & Streamlit)
- Built an end-to-end agentic application that screens resumes and prepares candidates for interviews.
- Created a 4-agent workflow using LangGraph to parse resumes, perform gap analysis, score job matching, and generate simulated interview questions.

Microservice Gateway (FastAPI & Redis)
- Developed a high-performance, rate-limited API gateway that handles 5,000+ requests per minute.
- Configured Redis caching to reduce database read overhead by 60%.
"""

def get_sample_job_description() -> str:
    return """Lead AI & Full-Stack Engineer
Company: FutureAI Solutions
Location: Remote (US) / Hybrid (San Francisco, CA)
Job Type: Full-time

ABOUT US
FutureAI Solutions is pioneering the next generation of enterprise automation. We build state-of-the-art agentic workflows and intelligent applications that empower teams to do their best work.

THE ROLE
We are seeking a Lead AI & Full-Stack Engineer to champion the development of our enterprise agentic platform. In this role, you will bridge the gap between frontend user experiences and complex backend agent systems, leading both the architectural design and hands-on coding.

KEY RESPONSIBILITIES
- Architect, build, and maintain production-ready agentic workflows using LangGraph, LangChain, and advanced LLMs (Gemini, Claude, GPT).
- Develop highly interactive, polished web interfaces in React and Next.js that interact with agentic backends via WebSockets and REST APIs.
- Design scalable backend APIs using FastAPI and coordinate complex database interactions with PostgreSQL and Redis.
- Lead system scaling and deployment using Kubernetes (EKS/GKE) and GCP (Google Cloud Platform).
- Mentor junior and mid-level engineers, enforcing clean code standards, writing test plans, and leading design reviews.

MINIMUM REQUIREMENTS
- 5+ years of software engineering experience in full-stack web applications.
- Strong proficiency in Python and TypeScript/React.
- Production experience deploying LLM applications using LangChain, LangGraph, and Vector Databases (FAISS, Pinecone, or Chroma).
- Deep experience with relational databases (PostgreSQL) and caching layers (Redis).
- Hands-on experience with containerization (Docker) and cloud deployment.

PREFERRED QUALIFICATIONS
- Experience with GCP (Google Cloud Platform) and Kubernetes (GKE).
- Prior experience in a tech lead or mentoring capacity.
- Contributions to open-source LLM frameworks or AI libraries.
"""
